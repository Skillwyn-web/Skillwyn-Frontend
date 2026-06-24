import io
import os
from typing import BinaryIO

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import docx
except Exception:
    docx = None

try:
    import pytesseract
    from PIL import Image
except Exception:
    pytesseract = None


def extract_text_from_pdf(pdf_input: BinaryIO | str) -> str:
    if not fitz:
        raise RuntimeError("PyMuPDF (fitz) is required for PDF parsing")

    full_text = ""
    if isinstance(pdf_input, str):
        pdf_doc = fitz.open(pdf_input)
    else:
        pdf_bytes = pdf_input.read()
        pdf_doc = fitz.open("pdf", pdf_bytes)

    for page in pdf_doc:
        full_text += page.get_text()

    pdf_doc.close()

    if not full_text.strip() and pytesseract:
        # Attempt OCR per page
        # Reopen as bytes
        if isinstance(pdf_input, str):
            from pdf2image import convert_from_path
            images = convert_from_path(pdf_input)
        else:
            from pdf2image import convert_from_bytes
            pdf_input.seek(0)
            images = convert_from_bytes(pdf_input.read())

        for img in images:
            full_text += pytesseract.image_to_string(img)

    if not full_text.strip():
        raise ValueError("No text found in PDF")

    return full_text


def extract_text_from_docx(file_input: BinaryIO | str) -> str:
    if not docx:
        raise RuntimeError("python-docx is required for DOCX parsing")

    if isinstance(file_input, str):
        document = docx.Document(file_input)
    else:
        file_input.seek(0)
        document = docx.Document(io.BytesIO(file_input.read()))

    full_text = "\n".join(p.text for p in document.paragraphs)
    if not full_text.strip():
        raise ValueError("No text found in DOCX")
    return full_text


def parse_resume(uploaded_file) -> str:
    """Detect file type and extract raw text. uploaded_file may be Starlette UploadFile or similar."""
    filename = getattr(uploaded_file, "filename", None) or str(uploaded_file)
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    if lower.endswith(".docx") or lower.endswith(".doc"):
        return extract_text_from_docx(uploaded_file)

    # Fallback: try reading as bytes and OCR if possible
    if pytesseract:
        try:
            if hasattr(uploaded_file, "read"):
                uploaded_file.seek(0)
                from PIL import Image
                img = Image.open(uploaded_file)
                return pytesseract.image_to_string(img)
        except Exception:
            pass

    raise ValueError("Unsupported file type or no text extracted")
